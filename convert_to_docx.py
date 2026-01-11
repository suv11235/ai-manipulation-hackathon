#!/usr/bin/env python3
"""Convert markdown report to Word document."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def parse_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to Word document."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Title (first #)
        if i == 0 and line.startswith('# '):
            title = line[2:].strip()
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # Headings
        if line.startswith('## '):
            heading = line[3:].strip()
            if heading == 'Authors':
                doc.add_heading(heading, level=1)
            elif heading in ['Abstract', 'Introduction', 'Methods', 'Results', 'Discussion and Conclusion', 'References', 'Appendix']:
                doc.add_heading(heading, level=1)
            else:
                doc.add_heading(heading, level=2)
            i += 1
            continue
        
        if line.startswith('### '):
            heading = line[4:].strip()
            doc.add_heading(heading, level=3)
            i += 1
            continue
        
        if line.startswith('#### '):
            heading = line[5:].strip()
            doc.add_heading(heading, level=4)
            i += 1
            continue
        
        # Horizontal rule
        if line.strip() == '---':
            i += 1
            continue
        
        # Tables (markdown table format)
        if '|' in line and i < len(lines) - 1:
            # Check if next line is a separator
            if i + 1 < len(lines) and '|' in lines[i + 1] and '---' in lines[i + 1]:
                # This is a table
                table_lines = []
                j = i
                while j < len(lines) and '|' in lines[j]:
                    if '---' not in lines[j]:
                        table_lines.append(lines[j].rstrip())
                    j += 1
                
                if table_lines:
                    # Parse table
                    headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
                    rows = []
                    for row_line in table_lines[1:]:
                        row = [cell.strip() for cell in row_line.split('|')[1:-1]]
                        rows.append(row)
                    
                    # Create Word table
                    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
                    table.style = 'Light Grid Accent 1'
                    
                    # Add headers
                    header_cells = table.rows[0].cells
                    for idx, header in enumerate(headers):
                        header_cells[idx].text = header
                        header_cells[idx].paragraphs[0].runs[0].font.bold = True
                    
                    # Add rows
                    for row_idx, row_data in enumerate(rows):
                        row_cells = table.rows[row_idx + 1].cells
                        for col_idx, cell_data in enumerate(row_data):
                            if col_idx < len(row_cells):
                                row_cells[col_idx].text = cell_data
                    
                    i = j
                    continue
        
        # Bold text (**text**)
        if '**' in line:
            para = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                else:
                    para.add_run(part)
            i += 1
            continue
        
        # Regular paragraph
        if line.strip():
            # Handle keywords line
            if line.strip().startswith('**Keywords**'):
                para = doc.add_paragraph()
                run = para.add_run('Keywords: ')
                run.bold = True
                keywords = line.split(':', 1)[1].strip()
                para.add_run(keywords)
            else:
                doc.add_paragraph(line.strip())
        else:
            # Empty line - add spacing
            doc.add_paragraph()
        
        i += 1
    
    doc.save(docx_file)
    print(f"Successfully converted {md_file} to {docx_file}")

if __name__ == '__main__':
    parse_markdown_to_docx('SUBMISSION_REPORT.md', 'SUBMISSION_REPORT.docx')

